from aiogram import Bot, Dispatcher, executor, types
from aiogram.types import callback_query
from aiogram.types import InputFile
from dotenv import load_dotenv
from aiogram.contrib.fsm_storage.memory import MemoryStorage
from aiogram.dispatcher.filters.state import State, StatesGroup
from aiogram.dispatcher import FSMContext
from datetime import date
from peewee import fn
from app import keyboards as kb
from app import database as db
from app import models as md
from docx import Document
import os
import pandas as pd
import yadisk
import asyncio

# Создание объекта хранилища памяти для сохранения состояний бота и данных сеансов
storage = MemoryStorage()

# Загрузка переменных среды из файла .env
load_dotenv()

# Инициализация бота
bot = Bot(os.getenv('TOKEN'))

# Инициализация диспетчера с передачей объекта бота и хранилища для работы с обработчиками и состояниями
dp = Dispatcher(bot=bot, storage=storage)

# Путь к папке с фотографиями
photo_directory = "C:/Users/Akelk/arts_store2/pythonProject/photo/"

# Запуск бд
async def on_startup(_):
    await db.db_start()
    print('Бот успешно запустился')

# Машина состояний для добавления товара
class NewOrder(StatesGroup):
    type = State()
    name = State()
    description = State()
    price = State()
    photo = State()

# Машина состояний для удаления товара
class DeleteBy(StatesGroup):
    name = State()
    id = State()

# обработчик команды /start
@dp.message_handler(commands=['start'])
async def cmd_start(message: types.Message):
    await db.cmd_start_db(message.from_user.id)
    await message.answer_sticker('CAACAgIAAxkBAAMpZBAAAfUO9xqQuhom1S8wBMW98ausAAI4CwACTuSZSzKxR9LZT4zQLwQ')
    await message.answer(f'{message.from_user.first_name}, добро пожаловать в магазин картин!',
                         reply_markup=kb.main)
    if message.from_user.id == int(os.getenv('ADMIN_ID')):
        await message.answer(f'Вы авторизовались как администратор!', reply_markup=kb.main_admin)

# Обработчик перехода в главное меню
@dp.message_handler(text='В главное меню')
async def process_main_menu(message: types.Message):
    if message.from_user.id == int(os.getenv('ADMIN_ID')):
        await message.answer('Выберите действие:', reply_markup=kb.main_admin)
    else:
        await message.answer('Выберите действие:', reply_markup=kb.main)

# Обработчик для каталога
@dp.message_handler(text='Каталог')
async def catalog(message: types.Message):
    await message.answer(f'Держи!', reply_markup=kb.catalog_list)


# Обработчик для кнопки "Удалить из корзины"
@dp.callback_query_handler(lambda c: c.data.startswith('delete_'))
async def delete_item_from_cart(call: types.CallbackQuery):
    user_id = call.from_user.id
    print(f"Содержимое callback_data: {call.data}")  # Вывод содержимого call.data для отладки
    item_id = int(call.data.split('_')[1])
    print(f"Извлеченный item_id: {item_id}")  # Вывод извлеченного item_id для отладки

    cart_item = md.Cart.get_or_none(item_id=item_id, user_id=user_id)  # Получить конкретный товар из корзины
    if cart_item:  # Проверить, найден ли товар
        cart_item.delete_instance()
        await bot.answer_callback_query(call.id, text="Товар удален из корзины")
    else:
        await bot.answer_callback_query(call.id, text="Товар не найден в корзине")

# test
@dp.message_handler(content_types=['document', 'photo'])
async def forward_message(message: types.Message):
    await bot.forward_message(os.getenv('GROUP_ID'), message.from_user.id, message.message_id)


# Обработчик для кнопки "Сделать заказ"
@dp.callback_query_handler(lambda c: c.data.startswith('order_'))
async def make_order(call: types.CallbackQuery):
    user_id = call.from_user.id
    username = f"https://t.me/{call.from_user.username}"

    print(f"Содержимое callback_data: {call.data}")
    item_id = int(call.data.split('_')[1])
    print(f"Извлеченный item_id: {item_id}")

    item = md.Item.get(md.Item.i_id == item_id)

    if item:
        photo = types.InputFile(item.photo)
        await bot.send_photo(chat_id=os.getenv('GROUP_ID'), photo=photo,
                         caption=f"Куплен товар item_id: {item_id}, пользователем {username}")
        await delete_photo_by_filename(item.photo)
        price = item.price

        # Создаем запись в таблице History
        history_entry = md.History.create(item_id=item_id, i_id=user_id, date=date.today(), price=price)
        history_entry.save()
        item.delete_instance()

        cart_item = md.Cart.get_or_none(item_id=item_id)  # Получить конкретный товар из корзины
        if cart_item:  # Проверить, найден ли товар
            cart_item.delete_instance()
            await bot.answer_callback_query(call.id, text="Заказ оформлен успешно")
        else:
            await bot.answer_callback_query(call.id, text="Вы не успели, товар куплен!")
    else:
        await bot.answer_callback_query(call.id, text="Вы не успели, товар куплен!")


# Функция для отображения корзины
@dp.message_handler(text='Корзина')
async def cart(message: types.Message):
    user_id = message.from_user.id

    # Проверяем наличие товаров в корзине перед удалением
    query = md.Item.select(md.Item.i_id)
    cart_items_to_delete = md.Cart.delete().where(~(md.Cart.item_id.in_(query)))
    deleted_rows = cart_items_to_delete.execute()

    subquery = (md.Cart
                .select(fn.MIN(md.Cart.id))
                .group_by(md.Cart.user_id, md.Cart.item_id))

    query = md.Cart.delete().where(md.Cart.id.not_in(subquery))
    query.execute()

    if deleted_rows:
        # Уведомляем пользователя об удаленных товарах из корзины
        await message.reply(f"Некоторые товары были удалены из вашей корзины, так как их больше нет в магазине.")

    # Получаем список товаров в корзине пользователя
    cart_items = (md.Item
                  .select(md.Item.i_id, md.Item.name, md.Item.description, md.Item.price, md.Item.photo)
                  .join(md.Cart, on=(md.Item.i_id == md.Cart.item_id))
                  .where(md.Cart.user_id == user_id))

    if cart_items:
        for cart_item in cart_items:
            item_id = cart_item.i_id
            name = cart_item.name
            description = cart_item.description
            price = cart_item.price
            photo = types.InputFile(cart_item.photo)

            # Подготавливаем сообщение о товаре
            caption = f"ID товара: {item_id}, Название: {name}, Описание: {description}, Цена: {price}"

            # Генерируем клавиатуру для товара в корзине
            keyboard = kb.generate_cart_keyboard(item_id)

            # Отправляем сообщение с изображением товара
            await bot.send_photo(chat_id=message.chat.id, photo=photo, caption=caption, reply_markup=keyboard)
    else:
        await message.reply("Ваша корзина пуста")


# Обработчик для контактов
@dp.message_handler(text='Контакты')
async def contacts(message: types.Message):
    await message.reply('@slonikmc')


# Обработчик для выгрузки данных
@dp.message_handler(text='Сделать выгрузку данных')
async def import_to(message: types.Message):
    if message.from_user.id == int(os.getenv('ADMIN_ID')):
        # Удаление файла, если уже существует
        if os.path.exists('history_output.xlsx'):
            os.remove('history_output.xlsx')
        if os.path.exists('history_output.docx'):
            os.remove('history_output.docx')

        # Выборка данных из таблицы History
        query = md.History.select()

        # Преобразование данных из выборки в DataFrame
        history_data = list(query.dicts())
        df = pd.DataFrame(history_data)

        # Экспорт DataFrame в Excel файл
        df.to_excel('history_output.xlsx', index=False)

        # Создание файла docx и запись данных из DataFrame построчно
        doc = Document()
        doc.add_heading('История Заказов', 0)

        for index, row in df.iterrows():
            for column in df.columns:
                doc.add_paragraph(f"{column}: {row[column]}")
            doc.add_paragraph('')

        doc.save('history_output.docx')

        # Отправка файлов пользователю
        with open('history_output.xlsx', 'rb') as file_xlsx, open('history_output.docx', 'rb') as file_docx:
            await message.reply_document(file_xlsx, caption="Excel файл")
            await message.reply_document(file_docx, caption="Word файл")

    else:
        await message.reply('У Вас недостаточно прав')

# Обработчик для истории заказов
@dp.message_handler(text='История заказов')
async def show_history(message: types.Message):
    if message.from_user.id == int(os.getenv('ADMIN_ID')):
        query = md.History.select()
    else:
        query = md.History.select().where(md.History.i_id == message.from_user.id)
    orders_info = "История заказов:\n"
    if query.count() > 0:
        for index, order in enumerate(query, start=1):
            order_info = f"Заказ №{index}: ID-товара - {order.item_id}, ID - {order.i_id}, Дата - {order.date}, Цена - {order.price}\n"
            orders_info += order_info
    else:
        orders_info = "Нет доступных заказов."

    await message.reply(orders_info)


# Обработчик входа в панель администратора
@dp.message_handler(text='Панель администратора')
async def contacts(message: types.Message):
    if message.from_user.id == int(os.getenv('ADMIN_ID')):
        await message.answer(f'Вы вошли в панель администратора', reply_markup=kb.admin_panel)
    else:
        await message.reply('Я тебя не понимаю.')

# Удаление фотографии по имени товара
async def delete_photo_by_filename(file_name):
    # Обработка пути к файлу
    file_path = os.path.join(photo_directory, file_name)
    if os.path.exists(file_path):
        os.remove(file_path)  # Удаляем файл
        return True
    else:
        return False

# Обработчик удаления товара
@dp.message_handler(text='Удалить товар')
async def delete_by_name(message: types.Message):
    if message.from_user.id == int(os.getenv('ADMIN_ID')):
        await message.answer(text=f"Пожалуйста, выберите товар который нужно удалить", reply_markup=kb.delete_panel)
    else:
        await message.reply('У вас нет прав для выполнения этой команды.')


# Обработчик для удаления по названию
@dp.message_handler(text='Удалить по названию')
async def delete_by_name(message: types.Message):
    if message.from_user.id == int(os.getenv('ADMIN_ID')):
        await message.answer('Введите название товара, который необходимо удалить')
        await DeleteBy.name.set()
    else:
        await message.reply('У вас нет прав для выполнения этой команды.')

# удаление по названию
@dp.message_handler(state=DeleteBy.name)
async def process_delete_by_name(message: types.Message, state: FSMContext):
    if message.from_user.id == int(os.getenv('ADMIN_ID')):
        name = message.text
        item = md.Item.get_or_none(name=name)
        if item:
            photo_name = item.photo
            item.delete_instance()
            await delete_photo_by_filename(photo_name)
            await state.finish()
            await message.answer(f'Товар с названием "{name}" успешно удален вместе с фотографией.')
        else:
            await message.answer('Товар с таким названием не найден.')
            await state.finish()
    else:
        await message.reply('У вас нет прав для выполнения этой команды.')

# Обработчик для удаления по id
@dp.message_handler(text='Удалить по id')
async def delete_by_id(message: types.Message):
    if message.from_user.id == int(os.getenv('ADMIN_ID')):
        await message.answer('Введите артикул товара, который необходимо удалить')
        await DeleteBy.id.set()
    else:
        await message.reply('У вас нет прав для выполнения этой команды.')

# удаление по id
@dp.message_handler(state=DeleteBy.id)
async def process_delete_by_id(message: types.Message, state: FSMContext):
    if message.from_user.id == int(os.getenv('ADMIN_ID')):
        id = message.text
        item = md.Item.get_or_none(i_id=id)  # Retrieve the item by name from the database
        if item:
            photo_name = item.photo  # Assuming the photo name is stored as an attribute on the Item model
            item.delete_instance()  # Delete the item from the database
            await delete_photo_by_filename(photo_name)  # Assuming a function to delete the photo by filename
            await state.finish()
            await message.answer(f'Товар с артикулом "{id}" успешно удален вместе с фотографией.')
        else:
            await message.answer('Товар с таким артикулом не найден.')
            await state.finish()
    else:
        await message.reply('У вас нет прав для выполнения этой команды.')

# добавление в корзину
@dp.callback_query_handler(lambda c: c.data.startswith('add_to_cart_'))
async def add_to_cart_handler(call: types.CallbackQuery):
    split_data = call.data.split('_')
    print(split_data)  # Добавляем эту строку для отображения значений, полученных после разделения
    item_id = int(split_data[3])
    user_id = call.from_user.id
    new_art = md.Cart.create(user_id=user_id, item_id=item_id)
    new_art.save()
    await call.message.answer("Товар добавлен в корзину!")


@dp.message_handler(text='Пейзажи')
async def handle_show_landscapes(message: types.Message):
    landscapes = md.Item.select().where(
        md.Item.category == 'peizaj')  # Assuming 'Пейзажи' is the category for landscapes
    if landscapes:
        for landscape in landscapes:
            photo_path = os.path.join(photo_directory,
                                      landscape.photo)  # Assuming 'photo' is the attribute representing the photo path
            with open(photo_path, 'rb') as photo_file:
                caption = f"Артикул: {landscape.i_id}, Название: {landscape.name}, Описание: {landscape.description}, Цена: {landscape.price}"
                keyboard = kb.generate_cart(landscape.i_id)  # Assuming a function to generate the cart keyboard
                await bot.send_photo(chat_id=message.chat.id, photo=types.InputFile(photo_file, filename='photo'),
                                     caption=caption, reply_markup=keyboard)
    else:
        await message.reply("Пейзажи не найдены в каталоге")


@dp.message_handler(text='Натюрморты')
async def handle_show_still_life(message: types.Message):
    still_life = md.Item.select().where(md.Item.category == 'naturmorts')
    if still_life:
        for still_l in still_life:
            photo_path = os.path.join(photo_directory,
                                      still_l.photo)  # Assuming 'photo' is the attribute representing the photo path
            with open(photo_path, 'rb') as photo_file:
                caption = f"Артикул: {still_l.i_id}, Название: {still_l.name}, Описание: {still_l.description}, Цена: {still_l.price}"
                keyboard = kb.generate_cart(still_l.i_id)  # Assuming a function to generate the cart keyboard
                await bot.send_photo(chat_id=message.chat.id, photo=types.InputFile(photo_file, filename='photo'),
                                     caption=caption, reply_markup=keyboard)
    else:
        await message.reply("Натюрморты не найдены в каталоге")


@dp.message_handler(text='Портреты')
async def handle_show_portrets(message: types.Message):
    portret = md.Item.select().where(md.Item.category == 'portrets')
    if portret:
        for prt in portret:
            photo_path = os.path.join(photo_directory,
                                      prt.photo)  # Assuming 'photo' is the attribute representing the photo path
            with open(photo_path, 'rb') as photo_file:
                caption = f"Артикул: {prt.i_id}, Название: {prt.name}, Описание: {prt.description}, Цена: {prt.price}"
                keyboard = kb.generate_cart(prt.i_id)  # Assuming a function to generate the cart keyboard
                await bot.send_photo(chat_id=message.chat.id, photo=types.InputFile(photo_file, filename='photo'),
                                     caption=caption, reply_markup=keyboard)
                await message.reply("Для заказа портрета перейдите в контакты")
    else:
        await message.reply("Портреты не найдены в каталоге")


# Обработчик для бэкапа
@dp.message_handler(text='Выполнить backup')
async def handle_backup_command(message: types.Message):
    # Вызов функции создания резервной копии с передачей объекта сообщения
    await db.upload_backup()
    await message.answer("БЭКАП ВЫПОЛНЕН")


# Обработка отмены
@dp.message_handler(text='Отмена', state='*')
async def cancel_add_item(message: types.Message, state: FSMContext):
    await state.finish()
    await message.answer('Добавление товара отменено.', reply_markup=kb.admin_panel)


# Добавление товара
@dp.message_handler(text='Добавить товар')
async def add_item_type(message: types.Message):
    if message.from_user.id == int(os.getenv('ADMIN_ID')):
        await NewOrder.type.set()
        await message.answer(f'Выберите тип товара', reply_markup=kb.catalog_list)
    else:
        await message.reply('Доступ запрещен.')


# Функция для обработки типа товара и перехода в новое состояние
@dp.callback_query_handler(state=NewOrder.type)
async def add_item_type(call: types.CallbackQuery, state: FSMContext):
    async with state.proxy() as data:
        data['type'] = call.data
    await call.message.answer(f'Напишите название', reply_markup=kb.cancel)
    await NewOrder.next()


@dp.message_handler(state=NewOrder.name)
async def add_item_name(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['name'] = message.text
    await message.answer('Напишите описание')
    await NewOrder.next()


@dp.message_handler(state=NewOrder.description)
async def add_item_name(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['description'] = message.text
    await message.answer('Введите цену')
    await NewOrder.next()


@dp.message_handler(state=NewOrder.price)
async def add_item_name(message: types.Message, state: FSMContext):
    try:
        price = int(message.text)  # Пытаемся преобразовать введенный текст в целое число
        # В случае успешного преобразования сохраняем цену в данных состояния
        async with state.proxy() as data:
            data['price'] = price
        await message.answer('Пожалуйста, прикрепите фотографию товара')
        await NewOrder.next()
    except ValueError:
        # В случае, если введенный текст не может быть преобразован в целое число
        await message.answer('Пожалуйста, введите целочисленную цену товара')


@dp.message_handler(content_types=['photo'], state=NewOrder.photo)
async def add_item_photo(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        # Save the photo to the server and store the file path in the data
        file_name = f"photo_{message.photo[-1].file_id}.jpg"
        photo_path = os.path.join(photo_directory, file_name)
        file_info = await bot.get_file(message.photo[-1].file_id)
        downloaded_file = await bot.download_file(file_info.file_path)
        with open(photo_path, 'wb') as new_file:
            new_file.write(downloaded_file.read())
        data['photo'] = photo_directory + file_name
        print(data['type'], data['name'], data['description'], data['price'], data['photo'])
        # Create a new item in the database using the provided details
        new_item = md.Item(category=data['type'], name=data['name'], description=data['description'],
                           price=data['price'], photo=data['photo'])
        new_item.save()  # Assuming 'save' is the method to save a new item to the database
    await message.answer('Товар успешно загружен!', reply_markup=kb.admin_panel)
    await state.finish()


# test
@dp.message_handler(content_types=['sticker'])
async def check_sticker(message: types.Message):
    await message.answer(message.sticker.file_id)
    await bot.send_message(message.from_user.id, message.chat.id)


# обработчик выбора какую категорию будем смотреть
@dp.callback_query_handler()
async def callback_query_keyboard(callback_query: types.CallbackQuery):
    if callback_query.data == 'naturmorts':
        await bot.send_message(chat_id=callback_query.from_user.id, text='Вы выбрали натюрморты')
        await handle_show_still_life(callback_query.message)
    elif callback_query.data == 'peizaj':
        await bot.send_message(chat_id=callback_query.from_user.id, text='Вы выбрали пейзажи')
        await handle_show_landscapes(callback_query.message)
    elif callback_query.data == 'portrets':
        await bot.send_message(chat_id=callback_query.from_user.id, text='Вы выбрали портреты')
        await handle_show_portrets(callback_query.message)


# В любой непонятной ситуации:
@dp.message_handler()
async def answer(message: types.Message):
    await message.reply('Я Вас не понимаю :(')


if __name__ == '__main__':
    executor.start_polling(dp, on_startup=on_startup, skip_updates=True)
